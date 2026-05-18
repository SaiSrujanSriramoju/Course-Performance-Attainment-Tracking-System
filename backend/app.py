from datetime import datetime
from functools import wraps
from io import BytesIO
import math
import os
import re

from flask import Flask, jsonify, request, send_file, send_from_directory, session
from flask_cors import CORS
from werkzeug.utils import secure_filename

import pandas as pd

from db import get_db_connection, init_db

# Flask API and UI server for the course outcome system.

app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("FLASK_SECRET_KEY", "dev-secret-key")
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"

# Allow the frontend pages to call the API during local development.
CORS(
    app,
    resources={r"/api/*": {"origins": ["http://127.0.0.1:5500", "http://localhost:5500"]}},
    supports_credentials=True,
)

# File upload settings for syllabus files.
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "..", "uploads", "syllabus")
ALLOWED_SYLLABUS_EXTENSIONS = {".pdf", ".doc", ".docx"}
MAX_SYLLABUS_SIZE_BYTES = 5 * 1024 * 1024

# Frontend UI directories for static page serving.
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")
ADMIN_UI_DIR = os.path.join(FRONTEND_DIR, "admin")
FACULTY_UI_DIR = os.path.join(FRONTEND_DIR, "faculty")
LOGIN_UI_DIR = os.path.join(FRONTEND_DIR, "login")
COURSES_UI_DIR = os.path.join(FRONTEND_DIR, "courses")
REPORTS_UI_DIR = os.path.join(FRONTEND_DIR, "reports")

app.config["MAX_CONTENT_LENGTH"] = MAX_SYLLABUS_SIZE_BYTES


# Check if the syllabus file has an allowed extension.
def allowed_syllabus_file(filename):
    _, ext = os.path.splitext(filename)
    return ext.lower() in ALLOWED_SYLLABUS_EXTENSIONS


# Clean PDF table cell values into a consistent string.
def _normalize_pdf_cell(value):
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


# Extract a PO number from a cell label.
def _parse_po_number(label):
    if not label:
        return None
    match = re.search(r"po\s*-?\s*(\d+)", label, re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1))


# Extract a CO number from a cell label.
def _parse_co_number(label):
    if not label:
        return None
    match = re.search(r"co\s*-?\s*(\d+)", label, re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1))


# Find all CO numbers mentioned inside a block of text.
def _extract_co_numbers_from_text(text):
    if not text:
        return []
    hits = re.findall(r"\bco\s*-?\s*(\d+)\b", text, flags=re.IGNORECASE)
    hits += re.findall(r"\bcourse\s+outcome\s*(\d+)\b", text, flags=re.IGNORECASE)
    values = []
    for item in hits:
        try:
            values.append(int(item))
        except (TypeError, ValueError):
            continue
    return values


# Parse the CO-PO matrix from a syllabus PDF (first usable table only).
def parse_co_po_matrix_from_pdf(file_path):
    try:
        import pdfplumber
    except Exception as exc:
        raise RuntimeError("pdfplumber not installed") from exc

    matrices = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            for table in tables:
                if not table:
                    continue
                normalized = [[_normalize_pdf_cell(cell) for cell in row] for row in table]

                header_row = None
                header_idx = None
                for idx, row in enumerate(normalized):
                    po_hits = [c for c in row if _parse_po_number(c) is not None]
                    if len(po_hits) >= 6:
                        header_row = row
                        header_idx = idx
                        break

                if header_row is None:
                    continue

                po_numbers = []
                for cell in header_row:
                    po_num = _parse_po_number(cell)
                    if po_num is not None:
                        po_numbers.append(po_num)

                if not po_numbers:
                    continue

                data_rows = normalized[header_idx + 1 :]
                matrix = {}
                for row in data_rows:
                    if not row:
                        continue
                    co_num = _parse_co_number(row[0])
                    if co_num is None:
                        continue
                    values = []
                    for cell in row[1 : 1 + len(po_numbers)]:
                        try:
                            values.append(int(str(cell).strip()))
                        except (TypeError, ValueError):
                            values.append(0)
                    if values:
                        matrix[co_num] = values

                if matrix:
                    matrices.append((po_numbers, matrix))

    if not matrices:
        return None

    po_numbers, matrix = matrices[0]
    return {"po_numbers": po_numbers, "matrix": matrix}


# Count the highest CO number found in a PDF syllabus.
def extract_co_count_from_pdf(file_path):
    try:
        import pdfplumber
    except Exception:
        return None

    co_numbers = set()
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            co_numbers.update(_extract_co_numbers_from_text(text))
    if not co_numbers:
        return None
    return max(co_numbers)


# Count the highest CO number found in a DOCX syllabus.
def extract_co_count_from_docx(file_path):
    try:
        import docx
    except Exception:
        return None

    co_numbers = set()
    document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        co_numbers.update(_extract_co_numbers_from_text(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                co_numbers.update(_extract_co_numbers_from_text(cell.text))
    if not co_numbers:
        return None
    return max(co_numbers)


# Count the highest CO number found in a DOC syllabus.
def extract_co_count_from_doc(file_path):
    try:
        import textract
    except Exception:
        return None

    try:
        raw = textract.process(file_path)
    except Exception:
        return None

    text = raw.decode("utf-8", errors="ignore") if isinstance(raw, (bytes, bytearray)) else str(raw)
    co_numbers = _extract_co_numbers_from_text(text)
    if not co_numbers:
        return None
    return max(co_numbers)


# Pick the right parser based on the syllabus file extension.
def extract_co_count_from_syllabus_path(file_path):
    if not file_path or not os.path.isfile(file_path):
        return None

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == ".pdf":
        return extract_co_count_from_pdf(file_path)
    if ext == ".docx":
        return extract_co_count_from_docx(file_path)
    if ext == ".doc":
        return extract_co_count_from_doc(file_path)
    return None


# Store a parsed CO-PO matrix into the database.
def save_co_po_matrix(course_id, parsed):
    if not parsed:
        return False

    po_numbers = parsed.get("po_numbers") or []
    matrix = parsed.get("matrix") or {}
    if not po_numbers or not matrix:
        return False

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM co_po_matrix WHERE course_id=%s", (course_id,))
        rows = []
        for co_num, values in matrix.items():
            for idx, po_num in enumerate(po_numbers):
                value = values[idx] if idx < len(values) else 0
                rows.append((course_id, int(co_num), int(po_num), int(value)))
        if rows:
            cursor.executemany(
                "INSERT INTO co_po_matrix (course_id, co_number, po_number, value) VALUES (%s, %s, %s, %s)",
                rows,
            )
        conn.commit()
        return True
    finally:
        cursor.close()
        conn.close()


# Unified JSON response helper.
def json_response(payload, status=200):
    return jsonify(payload), status


# Safely read JSON from the request body.
def get_request_json():
    data = request.get_json(silent=True)
    if data is None:
        return {}, json_response({"success": False, "message": "Invalid JSON body"}, 400)
    return data, None


# Require a logged-in user (any role).
def login_required(handler):
    @wraps(handler)
    def wrapper(*args, **kwargs):
        if not session.get("role"):
            return json_response({"success": False, "message": "Unauthorized"}, 401)
        return handler(*args, **kwargs)

    return wrapper


# Require a specific user role.
def role_required(role):
    def decorator(handler):
        @wraps(handler)
        def wrapper(*args, **kwargs):
            if not session.get("role"):
                return json_response({"success": False, "message": "Unauthorized"}, 401)
            if session.get("role") != role:
                return json_response({"success": False, "message": "Forbidden"}, 403)
            return handler(*args, **kwargs)

        return wrapper

    return decorator


# Check if a faculty member is assigned to a course.
def faculty_assigned_to_course(course_id, faculty_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "SELECT 1 FROM Course_Faculty WHERE course_id=%s AND faculty_id=%s",
            (course_id, faculty_id),
        )
        return cursor.fetchone() is not None
    finally:
        cursor.close()
        conn.close()


# Determine CO numbers for a course from the most reliable source.
def get_course_co_numbers(course_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT number_of_cos FROM Courses WHERE course_id=%s", (course_id,))
        row = cursor.fetchone()
        if row and row.get("number_of_cos"):
            return list(range(1, int(row["number_of_cos"]) + 1))

        cursor.execute(
            "SELECT DISTINCT co_number FROM co_po_matrix WHERE course_id=%s ORDER BY co_number",
            (course_id,),
        )
        rows = cursor.fetchall()
        if rows:
            return [r["co_number"] for r in rows]

        cursor.execute(
            "SELECT co_number FROM Course_Outcomes WHERE course_id=%s ORDER BY co_number",
            (course_id,),
        )
        rows = cursor.fetchall()
        return [r["co_number"] for r in rows]
    finally:
        cursor.close()
        conn.close()


# Fetch saved attainment values for a specific exam table.
def fetch_attainment_by_co(cursor, table_name, course_id, co_numbers):
    allowed_tables = {
        "minor1_attainment",
        "minor2_attainment",
        "minor3_attainment",
        "major_attainment",
    }
    if table_name not in allowed_tables:
        raise ValueError("Invalid attainment table")

    if not co_numbers:
        return {}

    placeholders = ", ".join(["%s"] * len(co_numbers))
    query = (
        f"SELECT co_number, attainment "
        f"FROM {table_name} "
        f"WHERE course_id=%s AND co_number IN ({placeholders})"
    )
    cursor.execute(query, [course_id, *co_numbers])
    rows = cursor.fetchall()
    return {int(row["co_number"]): row.get("attainment") for row in rows}


def fetch_attainment_cos(cursor, table_name, course_id):
    allowed_tables = {
        "minor1_attainment",
        "minor2_attainment",
        "minor3_attainment",
        "major_attainment",
    }
    if table_name not in allowed_tables:
        raise ValueError("Invalid attainment table")

    cursor.execute(
        f"SELECT DISTINCT co_number FROM {table_name} WHERE course_id=%s ORDER BY co_number",
        (course_id,),
    )
    rows = cursor.fetchall()
    return [int(row["co_number"]) for row in rows]


# Map exam type labels to their database table names.
def get_exam_table(exam_type):
    mapping = {
        "minor1": "minor1_attainment",
        "minor2": "minor2_attainment",
        "minor3": "minor3_attainment",
        "major": "major_attainment",
    }
    return mapping.get(str(exam_type).lower())


# Map exam types to the CO columns they use in Excel exports.
def get_exam_co_columns(exam_type):
    mapping = {
        "minor1": [1, 2],
        "minor2": [3, 4],
        "minor3": [5, 6],
        "major": [1, 2, 3, 4, 5, 6],
    }
    return mapping.get(str(exam_type).strip().lower())


# Extract a CO value from a flexible key/value payload.
def extract_co_value(data, co_number):
    if not isinstance(data, dict):
        return None
    for key, value in data.items():
        if key is None:
            continue
        match = re.search(r"co\s*[-_]?\s*(\d+)", str(key), re.IGNORECASE)
        if not match:
            continue
        try:
            if int(match.group(1)) == int(co_number):
                return value
        except (TypeError, ValueError):
            continue
    return None


# Normalize marks (including AB and blanks) into a consistent cell value.
def normalize_mark_cell(value):
    if value is None:
        return ""
    text = str(value).strip()
    if text == "":
        return ""
    if text.lower() == "ab":
        return "AB"
    try:
        number = float(text)
    except (TypeError, ValueError):
        return ""
    if math.isnan(number) or math.isinf(number):
        return ""
    return int(number) if number.is_integer() else number


# Export marks data to an Excel file (faculty only).
@app.route("/api/export_excel", methods=["POST"])
@role_required("faculty")
def export_marks_excel():
    data, error = get_request_json()
    if error:
        return error

    exam_type = (data.get("exam_type") or "").strip().lower()
    rows = data.get("data")
    max_marks = data.get("max_marks")

    co_numbers = data.get("co_numbers")
    if isinstance(co_numbers, list) and co_numbers:
        normalized = []
        seen = set()
        for value in co_numbers:
            try:
                co_num = int(value)
            except (TypeError, ValueError):
                return json_response({"success": False, "message": "Invalid co_numbers"}, 400)
            if co_num <= 0 or co_num in seen:
                continue
            normalized.append(co_num)
            seen.add(co_num)
        co_numbers = normalized
    else:
        co_numbers = get_exam_co_columns(exam_type)
        if not co_numbers:
            return json_response({"success": False, "message": "Invalid exam_type"}, 400)

    if not co_numbers:
        return json_response({"success": False, "message": "No CO columns provided"}, 400)
    if not isinstance(rows, list) or not rows:
        return json_response({"success": False, "message": "data list required"}, 400)
    if not isinstance(max_marks, dict):
        return json_response({"success": False, "message": "max_marks object required"}, 400)

    headers = ["Roll No"] + [f"CO-{co}" for co in co_numbers] + ["Tot"]

    max_row = ["Max Marks"]
    max_total = 0
    for co in co_numbers:
        raw = extract_co_value(max_marks, co)
        normalized = normalize_mark_cell(raw)
        if normalized == "" or str(normalized).lower() == "ab":
            return json_response({"success": False, "message": f"Missing max marks for CO-{co}"}, 400)
        try:
            max_total += float(normalized)
        except (TypeError, ValueError):
            return json_response({"success": False, "message": f"Invalid max marks for CO-{co}"}, 400)
        max_row.append(normalized)
    max_row.append(int(max_total) if float(max_total).is_integer() else round(max_total, 2))

    student_rows = []
    for entry in rows:
        if not isinstance(entry, dict):
            continue
        roll = str(entry.get("roll") or "").strip()
        if not roll:
            continue
        row = [roll]
        total = 0
        for co in co_numbers:
            raw = extract_co_value(entry, co)
            normalized = normalize_mark_cell(raw)
            row.append(normalized)
            try:
                total += float(normalized) if normalized not in ("", "AB") else 0
            except (TypeError, ValueError):
                total += 0
        total_cell = int(total) if float(total).is_integer() else round(total, 2)
        row.append(total_cell)
        student_rows.append(row)

    if not student_rows:
        return json_response({"success": False, "message": "No valid student rows found"}, 400)

    df = pd.DataFrame([max_row] + student_rows, columns=headers)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Marks")

    output.seek(0)
    filename = f"{exam_type}_marks.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# Load saved attainment level ranges for a course and faculty member.
def load_attainment_ranges(course_id, faculty_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co_number, level_number, lower_limit, upper_limit "
            "FROM co_attainment_levels "
            "WHERE course_id=%s AND faculty_id=%s "
            "ORDER BY co_number, level_number",
            (course_id, faculty_id),
        )
        rows = cursor.fetchall()
        grouped = {}
        for row in rows:
            grouped.setdefault(row["co_number"], []).append(row)
        return grouped
    finally:
        cursor.close()
        conn.close()


# Validate that attainment ranges are within 0-100.
def validate_attainment_ranges(ranges):
    grouped = {}
    for entry in ranges:
        co_number = entry.get("co_number")
        level_number = entry.get("level_number")
        lower_limit = entry.get("lower_limit")
        upper_limit = entry.get("upper_limit")

        if co_number is None or level_number is None:
            return False, "Missing CO or level number"
        try:
            co_number = int(co_number)
            level_number = int(level_number)
            lower_limit = int(lower_limit)
            upper_limit = int(upper_limit)
        except (TypeError, ValueError):
            return False, "Invalid range values"

        grouped.setdefault(co_number, []).append({
            "level_number": level_number,
            "lower_limit": lower_limit,
            "upper_limit": upper_limit,
        })

    for co_number, levels in grouped.items():
        levels.sort(key=lambda x: x["level_number"])
        expected_level = 1
        for level in levels:
            if level["level_number"] != expected_level:
                return False, f"CO{co_number} levels must be continuous from 1"
            if level["lower_limit"] >= level["upper_limit"]:
                return False, f"CO{co_number} lower must be less than upper"
            if level["lower_limit"] < 0:
                return False, f"CO{co_number} lower must be >= 0"
            if level["upper_limit"] > 100:
                return False, f"CO{co_number} upper must be <= 100"
            expected_level += 1

        if levels and levels[-1]["upper_limit"] != 100:
            return False, f"CO{co_number} last upper must be 100"

    return True, None


# Compute attainment percentage and level per CO using thresholds.
def compute_attainment_values(course_id, faculty_id):
    co_numbers = get_course_co_numbers(course_id)
    ranges = load_attainment_ranges(course_id, faculty_id)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co_number, co.co_id, t.threshold_mark "
            "FROM Course_Outcomes co "
            "LEFT JOIN CO_Thresholds t ON t.co_id = co.co_id "
            "WHERE co.course_id=%s",
            (course_id,),
        )
        co_rows = cursor.fetchall()
        co_id_map = {row["co_number"]: row for row in co_rows}

        attainment = {}
        for co_number in co_numbers:
            co_info = co_id_map.get(co_number)
            threshold = co_info.get("threshold_mark") if co_info else None
            co_id = co_info.get("co_id") if co_info else None

            percentage = 0.0
            if co_id and threshold is not None:
                cursor.execute(
                    "SELECT COUNT(*) AS total, "
                    "SUM(CASE WHEN marks >= %s THEN 1 ELSE 0 END) AS above "
                    "FROM Student_Scores WHERE co_id=%s",
                    (threshold, co_id),
                )
                row = cursor.fetchone() or {}
                total = row.get("total") or 0
                above = row.get("above") or 0
                if total:
                    percentage = (above / total) * 100

            level_value = 0
            for entry in ranges.get(co_number, []):
                if entry["lower_limit"] <= percentage <= entry["upper_limit"]:
                    level_value = entry["level_number"]
                    break

            attainment[co_number] = {
                "percentage": round(percentage, 2),
                "level": level_value,
            }

        return attainment
    finally:
        cursor.close()
        conn.close()


# Save per-CO threshold marks (creates CO rows if missing).
def save_co_thresholds(course_id, thresholds):
    if not thresholds:
        return

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        co_ids = []
        for entry in thresholds:
            co_number = int(entry["co_number"])
            threshold_mark = int(entry["threshold_mark"])

            cursor.execute(
                "SELECT co_id FROM Course_Outcomes WHERE course_id=%s AND co_number=%s",
                (course_id, co_number),
            )
            row = cursor.fetchone()
            if row:
                co_id = row["co_id"]
            else:
                cursor.execute(
                    "INSERT INTO Course_Outcomes (course_id, co_number, description) VALUES (%s, %s, %s)",
                    (course_id, co_number, ""),
                )
                co_id = cursor.lastrowid
            co_ids.append((co_id, threshold_mark))

        if co_ids:
            cursor.executemany(
                "DELETE FROM CO_Thresholds WHERE co_id=%s",
                [(co_id,) for co_id, _ in co_ids],
            )
            cursor.executemany(
                "INSERT INTO CO_Thresholds (co_id, threshold_mark) VALUES (%s, %s)",
                co_ids,
            )
        conn.commit()
    finally:
        cursor.close()
        conn.close()


# Unified login for admin and faculty.
@app.route("/api/login", methods=["POST"])
def login():
    data, error = get_request_json()
    if error:
        return error

    user_id = data.get("user_id")
    password = data.get("password")
    if not user_id or not password:
        return json_response({"success": False, "message": "user_id and password required"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT admin_id, username FROM Admin WHERE username=%s AND password=%s",
            (user_id, password),
        )
        admin = cursor.fetchone()
        if admin:
            session.clear()
            session["role"] = "admin"
            session["admin_id"] = admin["admin_id"]
            return json_response({"status": "success", "role": "admin"})

        cursor.execute(
            "SELECT faculty_id, name, email, department, first_login "
            "FROM Faculty WHERE faculty_id=%s AND password=%s",
            (user_id, password),
        )
        faculty = cursor.fetchone()
        if faculty:
            session.clear()
            session["role"] = "faculty"
            session["faculty_id"] = faculty["faculty_id"]
            return json_response({"status": "success", "role": "faculty"})

        return json_response({"status": "error", "message": "Invalid credentials"}, 401)
    finally:
        cursor.close()
        conn.close()


# Check current session role.
@app.route("/api/session", methods=["GET"])
def session_status():
    role = session.get("role")
    if not role:
        return json_response({"authenticated": False}, 401)

    payload = {"authenticated": True, "role": role}
    if role == "admin":
        payload["admin_id"] = session.get("admin_id")
    if role == "faculty":
        payload["faculty_id"] = session.get("faculty_id")
    return json_response(payload)


# Clear the session and log out.
@app.route("/api/logout", methods=["POST"])
def logout():
    session.clear()
    return json_response({"success": True})


# Admin-only login endpoint.
@app.route("/api/login/admin", methods=["POST"])
def login_admin():
    data, error = get_request_json()
    if error:
        return error

    username = data.get("username")
    password = data.get("password")
    if not username or not password:
        return json_response({"success": False, "message": "username and password required"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT admin_id, username FROM Admin WHERE username=%s AND password=%s",
            (username, password),
        )
        admin = cursor.fetchone()
        if not admin:
            return json_response({"success": False, "message": "Invalid credentials"}, 401)
        session.clear()
        session["role"] = "admin"
        session["admin_id"] = admin["admin_id"]
        return json_response({"success": True, "admin": admin})
    finally:
        cursor.close()
        conn.close()


# Faculty-only login endpoint.
@app.route("/api/login/faculty", methods=["POST"])
def login_faculty():
    data, error = get_request_json()
    if error:
        return error

    faculty_id = data.get("faculty_id")
    password = data.get("password")
    if not faculty_id or not password:
        return json_response({"success": False, "message": "faculty_id and password required"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT faculty_id, name, email, department, first_login "
            "FROM Faculty WHERE faculty_id=%s AND password=%s",
            (faculty_id, password),
        )
        faculty = cursor.fetchone()
        if not faculty:
            return json_response({"success": False, "message": "Invalid credentials"}, 401)
        session.clear()
        session["role"] = "faculty"
        session["faculty_id"] = faculty["faculty_id"]
        return json_response({"success": True, "faculty": faculty})
    finally:
        cursor.close()
        conn.close()


# Change password for the current session role.
@app.route("/api/change-password", methods=["POST"])
@login_required
def change_password():
    data, error = get_request_json()
    if error:
        return error

    role = data.get("role")
    new_password = data.get("new_password")
    old_password = data.get("old_password")

    if role not in ["admin", "faculty"]:
        return json_response({"success": False, "message": "role must be admin or faculty"}, 400)
    if session.get("role") != role:
        return json_response({"success": False, "message": "Forbidden"}, 403)

    if not new_password or not old_password:
        return json_response({"success": False, "message": "old_password and new_password required"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        if role == "admin":
            username = data.get("username")
            if not username:
                return json_response({"success": False, "message": "username required"}, 400)
            cursor.execute(
                "UPDATE Admin SET password=%s WHERE username=%s AND password=%s",
                (new_password, username, old_password),
            )
        else:
            faculty_id = data.get("faculty_id")
            if not faculty_id:
                return json_response({"success": False, "message": "faculty_id required"}, 400)
            cursor.execute(
                "UPDATE Faculty SET password=%s, first_login=0 "
                "WHERE faculty_id=%s AND password=%s",
                (new_password, faculty_id, old_password),
            )

        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Invalid credentials"}, 401)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Create a faculty account (admin only).
@app.route("/api/faculty", methods=["POST"])
@role_required("admin")
def create_faculty():
    data, error = get_request_json()
    if error:
        return error

    required = ["faculty_id", "name", "email", "department", "password"]
    if not all(data.get(k) for k in required):
        return json_response({"success": False, "message": "Missing faculty fields"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO Faculty (faculty_id, name, email, department, password, first_login) "
            "VALUES (%s, %s, %s, %s, %s, %s)",
            (
                data["faculty_id"],
                data["name"],
                data["email"],
                data["department"],
                data["password"],
                data.get("first_login", 1),
            ),
        )
        conn.commit()
        return json_response({"success": True}, 201)
    finally:
        cursor.close()
        conn.close()


# List faculty accounts (admin only).
@app.route("/api/faculty", methods=["GET"])
@role_required("admin")
def list_faculty():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT faculty_id, name, email, department, first_login FROM Faculty")
        rows = cursor.fetchall()
        return json_response({"success": True, "faculty": rows})
    finally:
        cursor.close()
        conn.close()


# Update faculty profile details (admin only).
@app.route("/api/faculty/<faculty_id>", methods=["PUT"])
@role_required("admin")
def update_faculty(faculty_id):
    data, error = get_request_json()
    if error:
        return error

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "UPDATE Faculty SET name=%s, email=%s, department=%s "
            "WHERE faculty_id=%s",
            (
                data.get("name"),
                data.get("email"),
                data.get("department"),
                faculty_id,
            ),
        )
        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Faculty not found"}, 404)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Delete a faculty account (admin only).
@app.route("/api/faculty/<faculty_id>", methods=["DELETE"])
@role_required("admin")
def delete_faculty(faculty_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM Course_Faculty WHERE faculty_id=%s", (faculty_id,))
        cursor.execute("DELETE FROM Faculty WHERE faculty_id=%s", (faculty_id,))
        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Faculty not found"}, 404)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Create a course and upload its syllabus (admin only).
@app.route("/api/courses", methods=["POST"])
@role_required("admin")
def create_course():
    form = request.form
    file = request.files.get("syllabus_file")

    if request.content_length and request.content_length > MAX_SYLLABUS_SIZE_BYTES:
        return json_response({"success": False, "message": "Syllabus file exceeds 5MB limit"}, 413)

    required = ["course_code", "course_name", "semester", "batch_name", "credits", "passing_marks"]
    if not all(form.get(k) for k in required):
        return json_response({"success": False, "message": "Missing course fields"}, 400)

    if not file or not file.filename:
        return json_response({"success": False, "message": "Syllabus file is required"}, 400)

    if not allowed_syllabus_file(file.filename):
        return json_response({"success": False, "message": "Only PDF or DOC/DOCX files are allowed"}, 400)

    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    safe_name = secure_filename(file.filename)
    _, ext = os.path.splitext(safe_name)
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    stored_name = f"{form.get('course_code')}_{timestamp}{ext.lower()}"
    file_path = os.path.join(UPLOAD_FOLDER, stored_name)
    file.save(file_path)

    syllabus_path = os.path.join("uploads", "syllabus", stored_name).replace("\\", "/")

    number_of_cos = form.get("number_of_cos")
    number_of_cos = int(number_of_cos) if number_of_cos else None

    if not number_of_cos:
        if ext.lower() == ".pdf":
            try:
                number_of_cos = extract_co_count_from_pdf(file_path)
            except Exception:
                number_of_cos = None
        elif ext.lower() == ".docx":
            try:
                number_of_cos = extract_co_count_from_docx(file_path)
            except Exception:
                number_of_cos = None
        elif ext.lower() == ".doc":
            try:
                number_of_cos = extract_co_count_from_doc(file_path)
            except Exception:
                number_of_cos = None

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO Courses (course_code, course_name, semester, batch_name, credits, passing_marks, syllabus_path, number_of_cos) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s)",
            (
                form["course_code"],
                form["course_name"],
                form["semester"],
                form["batch_name"],
                form["credits"],
                form["passing_marks"],
                syllabus_path,
                number_of_cos,
            ),
        )
        conn.commit()
        course_id = cursor.lastrowid
        if ext.lower() == ".pdf" and course_id:
            try:
                parsed = parse_co_po_matrix_from_pdf(file_path)
                save_co_po_matrix(course_id, parsed)
            except Exception:
                pass
        return json_response({"success": True, "syllabus_path": syllabus_path}, 201)
    finally:
        cursor.close()
        conn.close()


# List all courses (admin only).
@app.route("/api/courses", methods=["GET"])
@role_required("admin")
def list_courses():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT course_id, course_code, course_name, semester, batch_name, credits, "
            "passing_marks, syllabus_path, number_of_cos "
            "FROM Courses"
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "courses": rows})
    finally:
        cursor.close()
        conn.close()


# Serve uploaded syllabus files.
@app.route("/uploads/syllabus/<path:filename>", methods=["GET"])
def get_syllabus(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)


# Serve a syllabus to assigned faculty.
@app.route("/get_syllabus/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_syllabus_for_course(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT syllabus_path FROM Courses WHERE course_id=%s",
            (course_id,),
        )
        row = cursor.fetchone()
        if not row or not row.get("syllabus_path"):
            return json_response({"success": False, "message": "Syllabus not found"}, 404)

        file_path = os.path.join(BASE_DIR, row["syllabus_path"])
        if not os.path.isfile(file_path):
            return json_response({"success": False, "message": "Syllabus file missing"}, 404)
        return send_file(file_path, as_attachment=False)
    finally:
        cursor.close()
        conn.close()


# Return or infer the number of COs from a syllabus file.
@app.route("/api/syllabus/<int:course_id>/co-count", methods=["GET"])
@role_required("faculty")
def get_syllabus_co_count(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT syllabus_path, number_of_cos FROM Courses WHERE course_id=%s",
            (course_id,),
        )
        row = cursor.fetchone() or {}
        stored_count = row.get("number_of_cos")
        if stored_count:
            return json_response({"success": True, "co_count": int(stored_count)})

        syllabus_path = row.get("syllabus_path")
        if not syllabus_path:
            return json_response({"success": True, "co_count": 0})

        file_path = os.path.join(BASE_DIR, syllabus_path)
        try:
            count = extract_co_count_from_syllabus_path(file_path)
        except Exception:
            count = None

        count_value = int(count) if count else 0
        if count_value:
            cursor.execute(
                "UPDATE Courses SET number_of_cos=%s WHERE course_id=%s",
                (count_value, course_id),
            )
            conn.commit()

        return json_response({"success": True, "co_count": count_value})
    finally:
        cursor.close()
        conn.close()


# Update course details (admin only).
@app.route("/api/courses/<int:course_id>", methods=["PUT"])
@role_required("admin")
def update_course(course_id):
    data, error = get_request_json()
    if error:
        return error

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "UPDATE Courses SET course_code=%s, course_name=%s, semester=%s, batch_name=%s, credits=%s, "
            "passing_marks=%s, number_of_cos=%s WHERE course_id=%s",
            (
                data.get("course_code"),
                data.get("course_name"),
                data.get("semester"),
                data.get("batch_name"),
                data.get("credits"),
                data.get("passing_marks"),
                data.get("number_of_cos"),
                course_id,
            ),
        )
        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Course not found"}, 404)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Delete a course and all related records (admin only).
@app.route("/api/courses/<int:course_id>", methods=["DELETE"])
@role_required("admin")
def delete_course(course_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM Course_Faculty WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM co_attainment WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM co_attainment_levels WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM co_po_matrix WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM minor1_attainment WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM minor2_attainment WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM minor3_attainment WHERE course_id=%s", (course_id,))
        cursor.execute("DELETE FROM major_attainment WHERE course_id=%s", (course_id,))

        cursor.execute(
            "DELETE FROM CO_Thresholds WHERE co_id IN (SELECT co_id FROM Course_Outcomes WHERE course_id=%s)",
            (course_id,),
        )
        cursor.execute(
            "DELETE FROM Student_Scores WHERE co_id IN (SELECT co_id FROM Course_Outcomes WHERE course_id=%s)",
            (course_id,),
        )
        cursor.execute(
            "DELETE FROM Attainment_Levels WHERE co_id IN (SELECT co_id FROM Course_Outcomes WHERE course_id=%s)",
            (course_id,),
        )
        cursor.execute("DELETE FROM Course_Outcomes WHERE course_id=%s", (course_id,))

        cursor.execute("DELETE FROM Courses WHERE course_id=%s", (course_id,))
        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Course not found"}, 404)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Assign a faculty member to a course (admin only).
@app.route("/api/assign", methods=["POST"])
@role_required("admin")
def assign_faculty():
    data, error = get_request_json()
    if error:
        return error

    course_id = data.get("course_id")
    faculty_id = data.get("faculty_id")
    if not course_id or not faculty_id:
        return json_response({"success": False, "message": "course_id and faculty_id required"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO Course_Faculty (course_id, faculty_id) VALUES (%s, %s)",
            (course_id, faculty_id),
        )
        conn.commit()
        return json_response({"success": True}, 201)
    finally:
        cursor.close()
        conn.close()


# List course-faculty assignments (admin only).
@app.route("/api/assignments", methods=["GET"])
@role_required("admin")
def list_assignments():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT cf.id, cf.course_id, cf.faculty_id, c.course_code, c.course_name, "
            "f.name AS faculty_name, f.department "
            "FROM Course_Faculty cf "
            "JOIN Courses c ON c.course_id = cf.course_id "
            "JOIN Faculty f ON f.faculty_id = cf.faculty_id"
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "assignments": rows})
    finally:
        cursor.close()
        conn.close()


# Remove a course-faculty assignment (admin only).
@app.route("/api/assignments/<int:assignment_id>", methods=["DELETE"])
@role_required("admin")
def delete_assignment(assignment_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM Course_Faculty WHERE id=%s", (assignment_id,))
        conn.commit()
        if cursor.rowcount == 0:
            return json_response({"success": False, "message": "Assignment not found"}, 404)
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# List courses assigned to the logged-in faculty.
@app.route("/api/faculty/courses", methods=["GET"])
@role_required("faculty")
def list_faculty_courses():
    faculty_id = session.get("faculty_id")
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT c.course_id, c.course_code, c.course_name, c.semester, "
            "c.batch_name, c.credits, c.passing_marks, c.syllabus_path, c.number_of_cos "
            "FROM Courses c "
            "JOIN Course_Faculty cf ON cf.course_id = c.course_id "
            "WHERE cf.faculty_id = %s",
            (faculty_id,),
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "courses": rows})
    finally:
        cursor.close()
        conn.close()


# Get saved CO attainment values for a faculty member.
@app.route("/api/co_attainment/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_co_attainment(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co_number, attainment_value "
            "FROM co_attainment "
            "WHERE course_id=%s AND faculty_id=%s",
            (course_id, faculty_id),
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "attainment": rows})
    finally:
        cursor.close()
        conn.close()


# Save CO attainment values (faculty only).
@app.route("/api/save_co_attainment", methods=["POST"])
@role_required("faculty")
def save_co_attainment():
    data, error = get_request_json()
    if error:
        return error

    course_id = data.get("course_id")
    entries = data.get("entries")
    faculty_id = session.get("faculty_id")

    if not course_id or not isinstance(entries, list):
        return json_response({"success": False, "message": "course_id and entries required"}, 400)
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    normalized = []
    for entry in entries:
        co_number = entry.get("co_number")
        value = entry.get("attainment_value")
        try:
            value = float(value)
        except (TypeError, ValueError):
            return json_response({"success": False, "message": "Invalid attainment value"}, 400)
        if co_number is None:
            return json_response({"success": False, "message": "Invalid CO number"}, 400)
        if value < 0 or value > 3:
            return json_response({"success": False, "message": "Attainment must be 0-3"}, 400)
        normalized.append((course_id, faculty_id, int(co_number), value))

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "DELETE FROM co_attainment WHERE course_id=%s AND faculty_id=%s",
            (course_id, faculty_id),
        )
        cursor.executemany(
            "INSERT INTO co_attainment (course_id, faculty_id, co_number, attainment_value) "
            "VALUES (%s, %s, %s, %s)",
            normalized,
        )
        conn.commit()
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Save attainment levels and thresholds for a course.
@app.route("/api/save_attainment_ranges", methods=["POST"])
@app.route("/save_attainment_ranges", methods=["POST"])
@role_required("faculty")
def save_attainment_ranges():
    data, error = get_request_json()
    if error:
        return error

    course_id = data.get("course_id")
    ranges = data.get("ranges")
    thresholds = data.get("thresholds")
    faculty_id = session.get("faculty_id")

    if not course_id or not isinstance(ranges, list):
        return json_response({"success": False, "message": "course_id and ranges required"}, 400)
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    valid, message = validate_attainment_ranges(ranges)
    if not valid:
        return json_response({"success": False, "message": message}, 400)

    normalized = []
    for entry in ranges:
        normalized.append((
            course_id,
            faculty_id,
            int(entry["co_number"]),
            int(entry["level_number"]),
            int(entry["lower_limit"]),
            int(entry["upper_limit"]),
        ))

    if thresholds is not None:
        try:
            threshold_entries = []
            for entry in thresholds:
                co_number = entry.get("co_number")
                threshold_mark = entry.get("threshold_mark")
                if co_number is None or threshold_mark is None:
                    return json_response({"success": False, "message": "Invalid threshold data"}, 400)
                threshold_entries.append({
                    "co_number": int(co_number),
                    "threshold_mark": int(threshold_mark),
                })
        except (TypeError, ValueError):
            return json_response({"success": False, "message": "Invalid threshold data"}, 400)
        save_co_thresholds(course_id, threshold_entries)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "DELETE FROM co_attainment_levels WHERE course_id=%s AND faculty_id=%s",
            (course_id, faculty_id),
        )
        cursor.executemany(
            "INSERT INTO co_attainment_levels "
            "(course_id, faculty_id, co_number, level_number, lower_limit, upper_limit) "
            "VALUES (%s, %s, %s, %s, %s, %s)",
            normalized,
        )
        conn.commit()
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Fetch attainment ranges for UI use.
@app.route("/api/get_attainment_ranges/<int:course_id>", methods=["GET"])
@app.route("/get_attainment_ranges/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_attainment_ranges(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co_number, level_number, lower_limit, upper_limit "
            "FROM co_attainment_levels "
            "WHERE course_id=%s AND faculty_id=%s "
            "ORDER BY co_number, level_number",
            (course_id, faculty_id),
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "ranges": rows})
    finally:
        cursor.close()
        conn.close()


# Compute and return direct PO attainment from CO-PO matrix.
@app.route("/api/co_po_direct/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_co_po_direct(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        co_numbers = get_course_co_numbers(course_id)

        cursor.execute(
            "SELECT co_number, po_number, value "
            "FROM co_po_matrix WHERE course_id=%s",
            (course_id,),
        )
        matrix_rows = cursor.fetchall()

        attainment_by_co = compute_attainment_values(course_id, faculty_id)
        attainment_map = {co: data["level"] for co, data in attainment_by_co.items()}

        po_numbers = sorted({int(row["po_number"]) for row in matrix_rows})
        if not po_numbers:
            return json_response({"success": False, "message": "Syllabus CO-PO matrix not uploaded"}, 404)

        matrix = {co: {po: 0 for po in po_numbers} for co in co_numbers}
        for row in matrix_rows:
            co_num = int(row["co_number"])
            po_num = int(row["po_number"])
            if co_num in matrix and po_num in matrix[co_num]:
                matrix[co_num][po_num] = int(row["value"])

        column_averages = {}
        for po in po_numbers:
            values = [matrix[co][po] for co in co_numbers if matrix[co][po] > 0]
            column_averages[po] = round(sum(values) / len(values), 2) if values else 0

        co_po_direct = []
        for co in co_numbers:
            co_attainment = attainment_map.get(co, 0.0)
            for po in po_numbers:
                weight = matrix[co][po]
                direct_val = round(co_attainment * weight, 2)
                co_po_direct.append({
                    "co_number": co,
                    "po_number": po,
                    "value": direct_val,
                })

        po_attainment = []
        for po in po_numbers:
            weighted_sum = 0.0
            total_weight = 0.0
            for co in co_numbers:
                weight = matrix[co][po]
                if weight <= 0:
                    continue
                total_weight += weight
                weighted_sum += attainment_map.get(co, 0.0) * weight
            po_att = round(weighted_sum / total_weight, 2) if total_weight else 0
            po_attainment.append({"po_number": po, "attainment": po_att})

        return json_response({
            "success": True,
            "co_numbers": co_numbers,
            "po_numbers": po_numbers,
            "matrix": matrix,
            "column_averages": column_averages,
            "co_attainment": attainment_by_co,
            "co_po_direct": co_po_direct,
            "po_attainment": po_attainment,
        })
    finally:
        cursor.close()
        conn.close()


# Return CO-PO matrix and column averages.
@app.route("/api/co_po_matrix/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_co_po_matrix(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co_number, po_number, value AS weight "
            "FROM co_po_matrix WHERE course_id=%s",
            (course_id,),
        )
        rows = cursor.fetchall()
        if not rows:
            return json_response({"success": False, "message": "Syllabus CO-PO matrix not uploaded"}, 404)

        co_count = max(int(row["co_number"]) for row in rows)
        po_count = max(int(row["po_number"]) for row in rows)

        matrix = [[0 for _ in range(po_count)] for _ in range(co_count)]
        for row in rows:
            co_idx = int(row["co_number"]) - 1
            po_idx = int(row["po_number"]) - 1
            if 0 <= co_idx < co_count and 0 <= po_idx < po_count:
                matrix[co_idx][po_idx] = int(row["weight"])

        averages = []
        for po_idx in range(po_count):
            total = sum(matrix[co_idx][po_idx] for co_idx in range(co_count))
            averages.append(round(total / co_count, 2) if co_count else 0)

        return json_response({
            "success": True,
            "co_count": co_count,
            "po_count": po_count,
            "matrix": matrix,
            "averages": averages,
        })
    finally:
        cursor.close()
        conn.close()


# Compute final attainment mapping (internal + external).
@app.route("/api/final_mapping/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_final_mapping(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    co_numbers = get_course_co_numbers(course_id)
    if not co_numbers:
        return json_response({"success": False, "message": "No COs found"}, 404)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        minor1_cos = fetch_attainment_cos(cursor, "minor1_attainment", course_id)
        minor2_cos = fetch_attainment_cos(cursor, "minor2_attainment", course_id)
        minor3_cos = fetch_attainment_cos(cursor, "minor3_attainment", course_id)

        minor1_raw = fetch_attainment_by_co(cursor, "minor1_attainment", course_id, minor1_cos)
        minor2_raw = fetch_attainment_by_co(cursor, "minor2_attainment", course_id, minor2_cos)
        minor3_raw = fetch_attainment_by_co(cursor, "minor3_attainment", course_id, minor3_cos)
        major_raw = fetch_attainment_by_co(cursor, "major_attainment", course_id, co_numbers)

        def normalize_values(raw_map, cos):
            normalized = {}
            for co in cos:
                try:
                    value = float(raw_map.get(co))
                except (TypeError, ValueError):
                    value = 0.0
                if not math.isfinite(value):
                    value = 0.0
                normalized[co] = value
            return normalized

        minor1_map = normalize_values(minor1_raw, minor1_cos)
        minor2_map = normalize_values(minor2_raw, minor2_cos)
        minor3_map = normalize_values(minor3_raw, minor3_cos)
        major_map = normalize_values(major_raw, co_numbers)

        internal_by_co = {co: 0.0 for co in co_numbers}
        internal_by_co.update(minor1_map)
        internal_by_co.update(minor2_map)
        internal_by_co.update(minor3_map)

        results = []
        for co in co_numbers:
            internal = internal_by_co.get(co, 0.0)
            external = major_map.get(co, 0.0)
            direct = round((internal * 0.4) + (external * 0.6), 2)
            results.append({
                "co": f"CO{co}",
                "internal": internal,
                "external": external,
                "direct": direct,
            })

        return jsonify(results), 200
    finally:
        cursor.close()
        conn.close()


# Save attainment values for a specific exam type.
@app.route("/api/save_exam_attainment", methods=["POST"])
@role_required("faculty")
def save_exam_attainment():
    data, error = get_request_json()
    if error:
        return error

    course_id = data.get("course_id")
    exam_type = data.get("exam_type")
    entries = data.get("entries")
    faculty_id = session.get("faculty_id")

    if not course_id or not exam_type or not isinstance(entries, list):
        return json_response({"success": False, "message": "course_id, exam_type, entries required"}, 400)
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)

    table_name = get_exam_table(exam_type)
    if not table_name:
        return json_response({"success": False, "message": "Invalid exam_type"}, 400)

    normalized = []
    for entry in entries:
        co_number = entry.get("co_number")
        value = entry.get("attainment")
        try:
            co_number = int(co_number)
            value = float(value)
        except (TypeError, ValueError):
            return json_response({"success": False, "message": "Invalid attainment value"}, 400)
        if value < 0 or value > 3:
            return json_response({"success": False, "message": "Attainment must be 0-3"}, 400)
        normalized.append((course_id, co_number, value))

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            f"DELETE FROM {table_name} WHERE course_id=%s",
            (course_id,),
        )
        cursor.executemany(
            f"INSERT INTO {table_name} (course_id, co_number, attainment) VALUES (%s, %s, %s)",
            normalized,
        )
        conn.commit()
        return json_response({"success": True})
    finally:
        cursor.close()
        conn.close()


# Serve login UI.
@app.route("/login", methods=["GET"])
def login_page():
    return send_from_directory(LOGIN_UI_DIR, "index.html")


# Serve login UI assets.
@app.route("/login/<path:filename>", methods=["GET"])
def serve_login_assets(filename):
    return send_from_directory(LOGIN_UI_DIR, filename)


# Serve admin UI assets.
@app.route("/admin/<path:filename>", methods=["GET"])
def serve_admin_assets(filename):
    return send_from_directory(ADMIN_UI_DIR, filename)


# Serve admin dashboard page.
@app.route("/admin/dashboard", methods=["GET"])
def serve_admin_dashboard():
    return send_from_directory(ADMIN_UI_DIR, "dashboard.html")


# Serve faculty UI assets.
@app.route("/faculty/<path:filename>", methods=["GET"])
def serve_faculty_assets(filename):
    return send_from_directory(FACULTY_UI_DIR, filename)


# Serve faculty dashboard page.
@app.route("/faculty/dashboard", methods=["GET"])
def serve_faculty_dashboard():
    return send_from_directory(FACULTY_UI_DIR, "dashboard.html")


# Serve course UI assets.
@app.route("/courses/<path:filename>", methods=["GET"])
def serve_courses_assets(filename):
    return send_from_directory(COURSES_UI_DIR, filename)


# Serve courses page.
@app.route("/courses", methods=["GET"])
def serve_courses_page():
    return send_from_directory(COURSES_UI_DIR, "courses.html")


# Serve reports UI assets.
@app.route("/reports/<path:filename>", methods=["GET"])
def serve_reports_assets(filename):
    return send_from_directory(REPORTS_UI_DIR, filename)


# Serve reports page.
@app.route("/reports", methods=["GET"])
def serve_reports_page():
    return send_from_directory(REPORTS_UI_DIR, "reports.html")


# Create a course outcome entry.
@app.route("/api/co", methods=["POST"])
def create_co():
    data, error = get_request_json()
    if error:
        return error

    required = ["course_id", "co_number", "description"]
    if not all(data.get(k) for k in required):
        return json_response({"success": False, "message": "Missing CO fields"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO Course_Outcomes (course_id, co_number, description) VALUES (%s, %s, %s)",
            (data["course_id"], data["co_number"], data["description"]),
        )
        conn.commit()
        return json_response({"success": True}, 201)
    finally:
        cursor.close()
        conn.close()


# Create a threshold for a specific CO.
@app.route("/api/threshold", methods=["POST"])
def create_threshold():
    data, error = get_request_json()
    if error:
        return error

    required = ["co_id", "threshold_mark"]
    if not all(data.get(k) for k in required):
        return json_response({"success": False, "message": "Missing threshold fields"}, 400)

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO CO_Thresholds (co_id, threshold_mark) VALUES (%s, %s)",
            (data["co_id"], data["threshold_mark"]),
        )
        conn.commit()
        return json_response({"success": True}, 201)
    finally:
        cursor.close()
        conn.close()


# Fetch course outcomes and thresholds for a course.
@app.route("/api/co/<int:course_id>", methods=["GET"])
@role_required("faculty")
def get_course_outcomes(course_id):
    faculty_id = session.get("faculty_id")
    if not faculty_assigned_to_course(course_id, faculty_id):
        return json_response({"success": False, "message": "Forbidden"}, 403)
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co.co_id, co.co_number, co.description, t.threshold_mark "
            "FROM Course_Outcomes co "
            "LEFT JOIN CO_Thresholds t ON t.co_id = co.co_id "
            "WHERE co.course_id=%s",
            (course_id,),
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "outcomes": rows})
    finally:
        cursor.close()
        conn.close()


# Compute attainment levels based on thresholds and store them.
@app.route("/api/calculate-attainment", methods=["POST"])
def calculate_attainment():
    data, error = get_request_json()
    if error:
        return error

    course_id = data.get("course_id")
    levels = data.get("levels") or [
        {"min_percentage": 80, "level": 3},
        {"min_percentage": 60, "level": 2},
        {"min_percentage": 40, "level": 1},
    ]

    if not course_id:
        return json_response({"success": False, "message": "course_id required"}, 400)

    levels = sorted(levels, key=lambda x: x.get("min_percentage", 0), reverse=True)

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT co.co_id, t.threshold_mark "
            "FROM Course_Outcomes co "
            "LEFT JOIN CO_Thresholds t ON t.co_id = co.co_id "
            "WHERE co.course_id=%s",
            (course_id,),
        )
        outcomes = cursor.fetchall()

        results = []
        for outcome in outcomes:
            co_id = outcome["co_id"]
            threshold = outcome.get("threshold_mark") or 0

            cursor.execute(
                "SELECT AVG(marks) AS avg_marks FROM Student_Scores WHERE co_id=%s",
                (co_id,),
            )
            avg_marks = cursor.fetchone()["avg_marks"] or 0

            percentage = 0
            if threshold:
                percentage = (avg_marks / threshold) * 100

            level = 0
            for rule in levels:
                if percentage >= rule.get("min_percentage", 0):
                    level = rule.get("level", 0)
                    break

            cursor.execute(
                "INSERT INTO Attainment_Levels (co_id, percentage, level, calculated_at) "
                "VALUES (%s, %s, %s, %s)",
                (co_id, percentage, level, datetime.utcnow()),
            )

            results.append({
                "co_id": co_id,
                "percentage": percentage,
                "level": level,
            })

        conn.commit()
        return json_response({"success": True, "attainment": results})
    finally:
        cursor.close()
        conn.close()


# Read stored attainment results for a course.
@app.route("/api/attainment/<int:course_id>", methods=["GET"])
def get_attainment(course_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT a.attainment_id, a.co_id, a.percentage, a.level, a.calculated_at "
            "FROM Attainment_Levels a "
            "JOIN Course_Outcomes co ON co.co_id = a.co_id "
            "WHERE co.course_id=%s "
            "ORDER BY a.calculated_at DESC",
            (course_id,),
        )
        rows = cursor.fetchall()
        return json_response({"success": True, "attainment": rows})
    finally:
        cursor.close()
        conn.close()


# Development entrypoint.
if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=5000, debug=True)
